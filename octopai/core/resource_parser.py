"""
Resource Parser

This module provides parsers for various file formats to convert them into
skill resources that can be used for skill creation and evolution.
"""

import os
import base64
from typing import Dict, Any, Optional, List
from dataclasses import dataclass
from enum import Enum


class ResourceType(Enum):
    """Types of resources that can be parsed"""
    PDF = "pdf"
    DOC = "doc"
    DOCX = "docx"
    EXCEL = "excel"
    IMAGE = "image"
    VIDEO = "video"
    TEXT = "text"
    MARKDOWN = "markdown"
    HTML = "html"
    UNKNOWN = "unknown"


@dataclass
class ParsedResource:
    """Represents a parsed resource ready for skill usage"""
    file_path: str
    resource_type: ResourceType
    text_content: str
    metadata: Dict[str, Any]
    images: List[bytes] = None
    raw_data: bytes = None
    
    def __post_init__(self):
        if self.images is None:
            self.images = []
    
    def to_skill_resource(self) -> str:
        """Convert parsed resource to skill resource format"""
        parts = []
        
        parts.append(f"# Resource: {os.path.basename(self.file_path)}")
        parts.append(f"Type: {self.resource_type.value}")
        parts.append("")
        
        if self.metadata:
            parts.append("## Metadata")
            for key, value in self.metadata.items():
                parts.append(f"- {key}: {value}")
            parts.append("")
        
        if self.text_content:
            parts.append("## Content")
            parts.append(self.text_content)
        
        if self.images:
            parts.append(f"\n## Images: {len(self.images)} image(s) extracted")
        
        return "\n".join(parts)
    
    def get_base64_images(self) -> List[str]:
        """Get images as base64 encoded strings"""
        return [base64.b64encode(img).decode('utf-8') for img in self.images]


class BaseParser:
    """Base class for all resource parsers"""
    
    def can_parse(self, file_path: str) -> bool:
        """Check if this parser can handle the file"""
        raise NotImplementedError
    
    def parse(self, file_path: str) -> ParsedResource:
        """Parse the file and return a ParsedResource"""
        raise NotImplementedError


class TextParser(BaseParser):
    """Parser for plain text files"""
    
    TEXT_EXTENSIONS = {'.txt', '.md', '.markdown', '.rst', '.csv', '.json', '.yaml', '.yml'}
    
    def can_parse(self, file_path: str) -> bool:
        ext = os.path.splitext(file_path)[1].lower()
        return ext in self.TEXT_EXTENSIONS
    
    def parse(self, file_path: str) -> ParsedResource:
        ext = os.path.splitext(file_path)[1].lower()
        
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            content = f.read()
        
        resource_type = ResourceType.TEXT
        if ext in {'.md', '.markdown'}:
            resource_type = ResourceType.MARKDOWN
        elif ext == '.csv':
            resource_type = ResourceType.TEXT
        
        return ParsedResource(
            file_path=file_path,
            resource_type=resource_type,
            text_content=content,
            metadata={
                'file_size': os.path.getsize(file_path),
                'extension': ext
            }
        )


class PDFParser(BaseParser):
    """Parser for PDF files"""
    
    def can_parse(self, file_path: str) -> bool:
        return file_path.lower().endswith('.pdf')
    
    def parse(self, file_path: str) -> ParsedResource:
        text_content = ""
        images = []
        metadata = {
            'file_size': os.path.getsize(file_path)
        }
        
        try:
            import PyPDF2
            
            with open(file_path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                
                metadata['num_pages'] = len(reader.pages)
                if reader.metadata:
                    metadata['title'] = getattr(reader.metadata, 'title', '')
                    metadata['author'] = getattr(reader.metadata, 'author', '')
                
                for page in reader.pages:
                    text_content += page.extract_text() + "\n\n"
        
        except ImportError:
            text_content = "[PyPDF2 not installed. Install with: pip install PyPDF2]"
        except Exception as e:
            text_content = f"[Error parsing PDF: {str(e)}]"
        
        return ParsedResource(
            file_path=file_path,
            resource_type=ResourceType.PDF,
            text_content=text_content,
            metadata=metadata,
            images=images
        )


class DOCParser(BaseParser):
    """Parser for DOC/DOCX files"""
    
    def can_parse(self, file_path: str) -> bool:
        ext = os.path.splitext(file_path)[1].lower()
        return ext in {'.doc', '.docx'}
    
    def parse(self, file_path: str) -> ParsedResource:
        text_content = ""
        metadata = {
            'file_size': os.path.getsize(file_path)
        }
        
        ext = os.path.splitext(file_path)[1].lower()
        
        try:
            if ext == '.docx':
                from docx import Document
                doc = Document(file_path)
                
                for para in doc.paragraphs:
                    text_content += para.text + "\n"
                
                for table in doc.tables:
                    for row in table.rows:
                        row_text = " | ".join(cell.text for cell in row.cells)
                        text_content += row_text + "\n"
                    text_content += "\n"
            
            else:
                text_content = "[.doc parsing requires additional setup. Consider converting to .docx first.]"
        
        except ImportError:
            text_content = "[python-docx not installed. Install with: pip install python-docx]"
        except Exception as e:
            text_content = f"[Error parsing document: {str(e)}]"
        
        return ParsedResource(
            file_path=file_path,
            resource_type=ResourceType.DOCX if ext == '.docx' else ResourceType.DOC,
            text_content=text_content,
            metadata=metadata
        )


class ExcelParser(BaseParser):
    """Parser for Excel files"""
    
    def can_parse(self, file_path: str) -> bool:
        ext = os.path.splitext(file_path)[1].lower()
        return ext in {'.xlsx', '.xls', '.csv'}
    
    def parse(self, file_path: str) -> ParsedResource:
        text_content = ""
        metadata = {
            'file_size': os.path.getsize(file_path)
        }
        
        ext = os.path.splitext(file_path)[1].lower()
        
        try:
            import pandas as pd
            
            if ext == '.csv':
                df = pd.read_csv(file_path)
                text_content = df.to_string(index=False)
                metadata['num_rows'] = len(df)
                metadata['num_columns'] = len(df.columns)
                metadata['columns'] = list(df.columns)
            
            else:
                xl = pd.ExcelFile(file_path)
                metadata['sheet_names'] = xl.sheet_names
                
                for sheet_name in xl.sheet_names:
                    df = pd.read_excel(file_path, sheet_name=sheet_name)
                    text_content += f"=== Sheet: {sheet_name} ===\n"
                    text_content += df.to_string(index=False)
                    text_content += "\n\n"
        
        except ImportError:
            text_content = "[pandas not installed. Install with: pip install pandas openpyxl]"
        except Exception as e:
            text_content = f"[Error parsing Excel: {str(e)}]"
        
        return ParsedResource(
            file_path=file_path,
            resource_type=ResourceType.EXCEL,
            text_content=text_content,
            metadata=metadata
        )


class ImageParser(BaseParser):
    """Parser for image files"""
    
    IMAGE_EXTENSIONS = {'.jpg', '.jpeg', '.png', '.gif', '.bmp', '.webp', '.tiff'}
    
    def can_parse(self, file_path: str) -> bool:
        ext = os.path.splitext(file_path)[1].lower()
        return ext in self.IMAGE_EXTENSIONS
    
    def parse(self, file_path: str) -> ParsedResource:
        metadata = {
            'file_size': os.path.getsize(file_path)
        }
        
        with open(file_path, 'rb') as f:
            raw_data = f.read()
        
        try:
            from PIL import Image
            with Image.open(file_path) as img:
                metadata['width'] = img.width
                metadata['height'] = img.height
                metadata['format'] = img.format
                metadata['mode'] = img.mode
        except ImportError:
            pass
        except Exception as e:
            metadata['error'] = str(e)
        
        return ParsedResource(
            file_path=file_path,
            resource_type=ResourceType.IMAGE,
            text_content=f"Image file: {os.path.basename(file_path)}\n"
                        f"Size: {metadata.get('width', 'unknown')}x{metadata.get('height', 'unknown')}",
            metadata=metadata,
            images=[raw_data],
            raw_data=raw_data
        )


class VideoParser(BaseParser):
    """Parser for video files"""
    
    VIDEO_EXTENSIONS = {'.mp4', '.avi', '.mov', '.mkv', '.webm', '.flv'}
    
    def can_parse(self, file_path: str) -> bool:
        ext = os.path.splitext(file_path)[1].lower()
        return ext in self.VIDEO_EXTENSIONS
    
    def parse(self, file_path: str) -> ParsedResource:
        metadata = {
            'file_size': os.path.getsize(file_path)
        }
        
        with open(file_path, 'rb') as f:
            raw_data = f.read()
        
        try:
            import cv2
            cap = cv2.VideoCapture(file_path)
            if cap.isOpened():
                metadata['fps'] = cap.get(cv2.CAP_PROP_FPS)
                metadata['frame_count'] = int(cap.get(cv2.CAP_PROP_FRAME_COUNT))
                metadata['width'] = int(cap.get(cv2.CAP_PROP_FRAME_WIDTH))
                metadata['height'] = int(cap.get(cv2.CAP_PROP_FRAME_HEIGHT))
                duration = metadata['frame_count'] / metadata['fps'] if metadata['fps'] > 0 else 0
                metadata['duration_seconds'] = duration
                cap.release()
        except ImportError:
            pass
        except Exception as e:
            metadata['error'] = str(e)
        
        return ParsedResource(
            file_path=file_path,
            resource_type=ResourceType.VIDEO,
            text_content=f"Video file: {os.path.basename(file_path)}\n"
                        f"Duration: {metadata.get('duration_seconds', 'unknown')} seconds\n"
                        f"Resolution: {metadata.get('width', 'unknown')}x{metadata.get('height', 'unknown')}",
            metadata=metadata,
            raw_data=raw_data
        )


class HTMLParser(BaseParser):
    """Parser for HTML files"""
    
    def can_parse(self, file_path: str) -> bool:
        ext = os.path.splitext(file_path)[1].lower()
        return ext in {'.html', '.htm'}
    
    def parse(self, file_path: str) -> ParsedResource:
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            html_content = f.read()
        
        text_content = html_content
        
        try:
            from bs4 import BeautifulSoup
            soup = BeautifulSoup(html_content, 'html.parser')
            
            for script in soup(["script", "style"]):
                script.decompose()
            
            text_content = soup.get_text(separator='\n', strip=True)
            
            title = soup.title.string if soup.title else ''
            
        except ImportError:
            title = ''
        except Exception as e:
            text_content = f"[Error parsing HTML: {str(e)}]"
            title = ''
        
        return ParsedResource(
            file_path=file_path,
            resource_type=ResourceType.HTML,
            text_content=text_content,
            metadata={
                'file_size': os.path.getsize(file_path),
                'title': title
            }
        )


class ResourceParser:
    """Main resource parser that dispatches to appropriate parsers"""
    
    def __init__(self):
        self.parsers = [
            TextParser(),
            PDFParser(),
            DOCParser(),
            ExcelParser(),
            ImageParser(),
            VideoParser(),
            HTMLParser()
        ]
    
    def parse(self, file_path: str) -> ParsedResource:
        """
        Parse a file using the appropriate parser
        
        Args:
            file_path: Path to the file to parse
            
        Returns:
            ParsedResource object
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        for parser in self.parsers:
            if parser.can_parse(file_path):
                return parser.parse(file_path)
        
        return ParsedResource(
            file_path=file_path,
            resource_type=ResourceType.UNKNOWN,
            text_content=f"Unknown file type: {file_path}",
            metadata={
                'file_size': os.path.getsize(file_path)
            }
        )
    
    def parse_to_skill_resource(self, file_path: str) -> str:
        """
        Parse a file and convert directly to skill resource format
        
        Args:
            file_path: Path to the file to parse
            
        Returns:
            String in skill resource format
        """
        parsed = self.parse(file_path)
        return parsed.to_skill_resource()


def parse_resource(file_path: str) -> ParsedResource:
    """
    Convenience function to parse a resource
    
    Args:
        file_path: Path to the file
        
    Returns:
        ParsedResource
    """
    parser = ResourceParser()
    return parser.parse(file_path)


def parse_to_skill_resource(file_path: str) -> str:
    """
    Convenience function to parse a file to skill resource format
    
    Args:
        file_path: Path to the file
        
    Returns:
        String in skill resource format
    """
    parser = ResourceParser()
    return parser.parse_to_skill_resource(file_path)
